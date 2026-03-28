import os
import shutil
import stat
import sys
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from langchain_chroma import Chroma
from langchain_core.documents import Document as LangchainDocument
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.entities import (
    ChatMessage,
    ChatSession,
    Document,
    DocumentContent,
    DocumentChunk,
    IndexStatus,
    MessageRole,
)


def ensure_storage_dir() -> Path:
    path = Path(settings.storage_dir)
    path.mkdir(parents=True, exist_ok=True)
    return path


def save_document_file(document_id: str, filename: str, data: bytes) -> str:
    storage = ensure_storage_dir()
    safe_name = filename.replace("/", "_")
    target = storage / f"{document_id}_upload_{uuid.uuid4().hex[:8]}_{safe_name}"
    target.write_bytes(data)
    return str(target)


def extract_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()
    if suffix == ".docx":
        doc = DocxDocument(str(path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    raise ValueError("Unsupported format; upload PDF or DOCX only.")


def chroma_dir_for_document(document_id: str) -> Path:
    if settings.chroma_persist_root:
        return Path(settings.chroma_persist_root) / document_id
    return ensure_storage_dir() / "chroma" / document_id


def latest_chroma_dir_for_document(document_id: str) -> Path:
    """Resolve the directory used for retrieval. Supports both legacy single-dir layout and new per-run subdirectories."""
    base = chroma_dir_for_document(document_id)
    if not base.exists():
        return base

    direct_db = base / "chroma.sqlite3"
    if direct_db.exists():
        return base

    run_dirs = [p for p in base.iterdir() if p.is_dir()]
    if not run_dirs:
        return base
    run_dirs.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return run_dirs[0]


def _chmod_u_rw(path: Path) -> None:
    try:
        mode = path.stat().st_mode
        os.chmod(path, mode | stat.S_IWRITE | stat.S_IREAD)
    except OSError:
        pass


def _rmtree_onexc(func, path: str, exc: BaseException) -> None:
    try:
        _chmod_u_rw(Path(path))
        func(path)
    except OSError:
        raise exc


def _rmtree_onerror(func, path: str, exc_info) -> None:
    try:
        _chmod_u_rw(Path(path))
        func(path)
    except OSError:
        pass


def clear_document_chroma(document_id: str) -> None:
    chroma_path = chroma_dir_for_document(document_id)
    if chroma_path.exists():
        if sys.version_info >= (3, 12):
            shutil.rmtree(chroma_path, onexc=_rmtree_onexc)
        else:
            shutil.rmtree(chroma_path, onerror=_rmtree_onerror)
    chroma_path.mkdir(parents=True, exist_ok=True)


def _assert_dir_writable(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)
    probe = path / ".write_probe_delete_me"
    try:
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
    except OSError as exc:
        raise RuntimeError(
            f"Directory is not writable: {path}. "
            "Fix ownership/permissions on document_agent_service/storage (or set DOCUMENT_AGENT_STORAGE_DIR "
            "to a folder your user owns, e.g. /tmp/doc_agent_data)."
        ) from exc


def create_or_update_document(
    db: Session,
    *,
    document_id: str,
    tenant_id: str,
    title: str,
    owner_id: str | None,
) -> Document:
    doc = db.get(Document, document_id)
    if not doc:
        doc = Document(id=document_id, tenant_id=tenant_id, title=title)
        db.add(doc)
    doc.owner_id = owner_id
    doc.index_status = IndexStatus.pending
    db.commit()
    db.refresh(doc)
    return doc


def split_into_chunks(text: str, chunk_size: int, overlap: int) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    return splitter.split_text(text)


def run_indexing_job(db: Session, document_id: str) -> int:
    doc = db.get(Document, document_id)
    if not doc:
        raise ValueError("document not found")
    content = db.execute(
        select(DocumentContent).where(DocumentContent.document_id == document_id).order_by(DocumentContent.id.desc())
    ).scalars().first()
    if not content:
        raise ValueError("document content not found")
    if not content.full_text.strip():
        raise ValueError("document is empty")

    doc.index_status = IndexStatus.processing
    db.commit()

    db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == document_id))
    db.commit()

    base_chroma_path = chroma_dir_for_document(document_id).resolve()
    base_chroma_path.mkdir(parents=True, exist_ok=True)
    _assert_dir_writable(base_chroma_path)

    chunk_texts = split_into_chunks(content.full_text, settings.chunk_size, settings.chunk_overlap)
    for idx, chunk in enumerate(chunk_texts):
        db.add(
            DocumentChunk(
                document_id=document_id,
                content_id=content.id,
                chunk_index=idx,
                text=chunk,
                vector_ref=f"chroma-{document_id}-{idx}",
            )
        )
    db.commit()

    rows = db.execute(
        select(DocumentChunk)
        .where(DocumentChunk.document_id == document_id)
        .order_by(DocumentChunk.chunk_index.asc())
    ).scalars().all()

    documents = [
        LangchainDocument(
            page_content=c.text,
            metadata={"chunk_id": c.id, "chunk_index": c.chunk_index},
        )
        for c in rows
    ]

    run_chroma_path = base_chroma_path / f"run_{uuid.uuid4().hex[:8]}"
    run_chroma_path.mkdir(parents=True, exist_ok=True)

    embedding = OpenAIEmbeddings(
        model=settings.openai_embed_model,
        api_key=settings.openai_api_key,
    )

    try:
        Chroma.from_documents(
            documents=documents,
            embedding=embedding,
            persist_directory=str(run_chroma_path),
        )
    except Exception:
        doc.index_status = IndexStatus.failed
        db.commit()
        raise

    doc.index_status = IndexStatus.ready
    db.commit()
    return len(rows)


def create_chat_session(db: Session, document_id: str, tenant_id: str, user_id: str) -> ChatSession:
    session = ChatSession(id=uuid.uuid4().hex, document_id=document_id, tenant_id=tenant_id, user_id=user_id)
    db.add(session)
    db.commit()
    db.refresh(session)
    return session


DOCUMENT_PROMPT = ChatPromptTemplate.from_template(
    """You are an assistant helping users understand document content.
Answer using only the context below. Be concise.
If the context does not contain enough information, say so clearly.

Context:
{context}

Question: {input}

Answer:"""
)


def answer_with_rag(db: Session, session_id: str, message: str, top_k: int) -> tuple[str, list[dict]]:
    session = db.get(ChatSession, session_id)
    if not session:
        raise ValueError("chat session not found")

    doc = db.get(Document, session.document_id)
    if not doc:
        raise ValueError("document not found")

    if doc.index_status != IndexStatus.ready:
        err = "Document is not indexed yet. Upload a PDF or DOCX, then run indexing, and wait until status is READY."
        db.add(ChatMessage(session_id=session_id, role=MessageRole.user, content=message))
        db.add(
            ChatMessage(
                session_id=session_id,
                role=MessageRole.assistant,
                content=err,
                citations_json={"items": []},
            )
        )
        db.commit()
        return err, []

    chroma_path = latest_chroma_dir_for_document(session.document_id).resolve()
    llm = ChatOpenAI(
        model=settings.openai_chat_model,
        api_key=settings.openai_api_key,
        temperature=0,
    )

    if not chroma_path.exists():
        err = "No search index on disk for this document. Run indexing again after uploading content."
        db.add(ChatMessage(session_id=session_id, role=MessageRole.user, content=message))
        db.add(
            ChatMessage(
                session_id=session_id,
                role=MessageRole.assistant,
                content=err,
                citations_json={"items": []},
            )
        )
        db.commit()
        return err, []

    embeddings = OpenAIEmbeddings(
        model=settings.openai_embed_model,
        api_key=settings.openai_api_key,
    )

    citations: list[dict] = []
    answer: str

    try:
        vector_store = Chroma(
            persist_directory=str(chroma_path.resolve()),
            embedding_function=embeddings,
        )
        docs = vector_store.similarity_search(message, k=top_k)
        citations = [
            {
                "chunk_id": d.metadata.get("chunk_id"),
                "chunk_index": d.metadata.get("chunk_index"),
            }
            for d in docs
        ]
        context = "\n\n".join(d.page_content for d in docs)
        chain = DOCUMENT_PROMPT | llm
        if not context.strip():
            result = chain.invoke({"context": "(No matching passages retrieved.)", "input": message})
        else:
            result = chain.invoke({"context": context, "input": message})
        answer = result.content if hasattr(result, "content") else str(result)
    except Exception as exc:
        try:
            chain = DOCUMENT_PROMPT | llm
            result = chain.invoke(
                {
                    "context": f"(Retrieval failed: {exc!s}. Answering without retrieved passages.)",
                    "input": message,
                }
            )
            answer = result.content if hasattr(result, "content") else str(result)
        except Exception as exc2:
            answer = (
                "The document assistant could not reach OpenAI. Ensure your API key is valid and has access to "
                f"model `{settings.openai_chat_model}`. "
                f"Details: {exc2!s}"
            )

    db.add(ChatMessage(session_id=session_id, role=MessageRole.user, content=message))
    db.add(
        ChatMessage(
            session_id=session_id,
            role=MessageRole.assistant,
            content=answer,
            citations_json={"items": citations},
        )
    )
    db.commit()
    return answer, citations


def generate_document_summary(db: Session, document_id: str) -> tuple[str | None, str | None, str | None]:
    """Generate a summary of the document and extract key sections (title, author, summary)."""
    doc = db.get(Document, document_id)
    if not doc:
        raise ValueError("Document not found")

    content = (
        db.query(DocumentContent)
        .filter(DocumentContent.document_id == document_id)
        .order_by(DocumentContent.id.desc())
        .first()
    )
    if not content or not content.full_text:
        return None, None, None

    text = content.full_text
    # Take first 3000 chars for summary to avoid token limits
    text_excerpt = text[:3000]

    llm = ChatOpenAI(
        model=settings.openai_chat_model,
        api_key=settings.openai_api_key,
        temperature=0,
    )

    prompt = ChatPromptTemplate.from_template(
        """Analyze this document excerpt and extract:
1. Document Title (just the title, or "Unknown" if not found)
2. Author (just the name, or "Unknown" if not found)
3. Brief Summary (2-3 sentences about the document)

Format your response exactly as:
TITLE: [title here]
AUTHOR: [author here]
SUMMARY: [summary here]

Document:
{text}"""
    )

    try:
        result = (prompt | llm).invoke({"text": text_excerpt})
        response_text = result.content if hasattr(result, "content") else str(result)

        # Parse the response
        title = "Unknown"
        author = "Unknown"
        summary = "Unable to generate summary"

        for line in response_text.split("\n"):
            if line.startswith("TITLE:"):
                title = line.replace("TITLE:", "").strip()
            elif line.startswith("AUTHOR:"):
                author = line.replace("AUTHOR:", "").strip()
            elif line.startswith("SUMMARY:"):
                summary = line.replace("SUMMARY:", "").strip()

        return title, author, summary
    except Exception as exc:
        return None, None, f"Error generating summary: {exc!s}"
